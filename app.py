from __future__ import annotations

import io
from pathlib import Path

import pandas as pd
import streamlit as st
from docx import Document

from matcher import Profile, parse_likes, top_matches


EXPECTED_COLUMNS = ["name", "gender", "likes", "age", "city", "bio", "photo_url"]
BASE_DIR = Path(__file__).resolve().parent
DUMMY_DATA_PATH = BASE_DIR / "data" / "dummy_profiles.csv"


def inject_styles() -> None:
    st.markdown(
        """
        <style>
            @import url('https://fonts.googleapis.com/css2?family=Fraunces:wght@500;700&family=Space+Grotesk:wght@400;500;700&display=swap');

            .stApp {
                background:
                    radial-gradient(circle at 5% 10%, rgba(255, 202, 213, 0.35), transparent 35%),
                    radial-gradient(circle at 90% 5%, rgba(182, 225, 252, 0.4), transparent 30%),
                    linear-gradient(180deg, #fffaf5 0%, #f6fbff 45%, #ffffff 100%);
                font-family: 'Space Grotesk', sans-serif;
            }

            h1, h2, h3 {
                font-family: 'Fraunces', serif;
                letter-spacing: 0.3px;
            }

            .hero {
                border-radius: 18px;
                padding: 22px;
                background: linear-gradient(120deg, #ff6f61 0%, #ff9f68 45%, #ffc77e 100%);
                color: #1f1f1f;
                box-shadow: 0 10px 28px rgba(0, 0, 0, 0.10);
                margin-bottom: 14px;
            }

            .stat-pill {
                border: 1px solid rgba(0, 0, 0, 0.08);
                border-radius: 999px;
                padding: 6px 12px;
                background: rgba(255, 255, 255, 0.85);
                display: inline-block;
                margin-right: 6px;
                margin-bottom: 6px;
                font-size: 0.84rem;
            }

            .profile-card {
                border: 1px solid rgba(0, 0, 0, 0.09);
                border-radius: 16px;
                background: #ffffffcc;
                padding: 16px;
                margin-bottom: 10px;
                box-shadow: 0 6px 18px rgba(34, 34, 34, 0.06);
            }

            .match-card {
                border: 1px solid rgba(0, 0, 0, 0.08);
                border-left: 8px solid #ff7b54;
                border-radius: 14px;
                padding: 14px;
                background: #fff;
                margin-bottom: 10px;
            }

            .match-score {
                font-size: 1.05rem;
                font-weight: 700;
            }

            .score-pill {
                background: #ffe9de;
                color: #812f17;
                border-radius: 999px;
                padding: 4px 10px;
                font-size: 0.82rem;
                font-weight: 700;
                border: 1px solid #ffd0be;
                white-space: nowrap;
            }

            .public-match-card {
                border: 1px solid rgba(0, 0, 0, 0.1);
                border-radius: 14px;
                background: #ffffff;
                padding: 10px;
                margin-bottom: 12px;
                box-shadow: 0 6px 18px rgba(34, 34, 34, 0.08);
            }

            .gallery-card {
                border: 1px solid rgba(0, 0, 0, 0.1);
                border-radius: 14px;
                background: #ffffff;
                padding: 8px 10px 12px 10px;
                box-shadow: 0 8px 20px rgba(31, 41, 55, 0.08);
                min-height: 280px;
            }

            .gallery-title {
                margin-top: 6px;
                font-weight: 700;
                font-size: 1rem;
            }

            .gallery-meta {
                margin-top: 2px;
                color: #4b5563;
                font-size: 0.83rem;
            }

            .gallery-bio {
                margin-top: 6px;
                font-size: 0.87rem;
                color: #1f2937;
                line-height: 1.35;
            }

            /* Keep Streamlit input widgets visually explicit. */
            [data-testid="stTextInput"] input,
            [data-testid="stTextArea"] textarea,
            [data-testid="stSelectbox"] [data-baseweb="select"] > div {
                background: #ffffff !important;
                color: #1f2937 !important;
                border: 1px solid #c9d6e2 !important;
                border-radius: 12px !important;
            }

            [data-testid="stTextInput"] label p,
            [data-testid="stTextArea"] label p,
            [data-testid="stSelectbox"] label p,
            [data-testid="stSlider"] label p,
            [data-testid="stNumberInput"] label p {
                color: #213547 !important;
                font-weight: 600 !important;
            }
        </style>
        """,
        unsafe_allow_html=True,
    )


def profile_likes_html(likes: set[str]) -> str:
    if not likes:
        return '<span class="stat-pill">No likes listed</span>'
    return "".join(f'<span class="stat-pill">{item}</span>' for item in sorted(likes))


def pick_profile_by_name(profiles: list[Profile], selected_name: str) -> Profile | None:
    for profile in profiles:
        if profile.name == selected_name:
            return profile
    return None


def profile_exists(profiles: list[Profile], name: str) -> bool:
    target = name.strip().lower()
    return any(p.name.strip().lower() == target for p in profiles)


def profile_exists_except(profiles: list[Profile], name: str, exclude_name: str) -> bool:
    target = name.strip().lower()
    excluded = exclude_name.strip().lower()
    return any(
        p.name.strip().lower() == target and p.name.strip().lower() != excluded
        for p in profiles
    )


def reset_profiles_in_memory(base_profiles: list[Profile], source_key: str) -> None:
    existing_key = st.session_state.get("source_key")
    if existing_key != source_key:
        st.session_state["source_key"] = source_key
        st.session_state["base_profiles"] = base_profiles

    if "base_profiles" not in st.session_state:
        st.session_state["base_profiles"] = base_profiles
    if "public_submissions" not in st.session_state:
        st.session_state["public_submissions"] = []


def combined_profiles() -> list[Profile]:
    return list(st.session_state.get("base_profiles", [])) + list(st.session_state.get("public_submissions", []))


def normalize_gender(value: str | None) -> str:
    raw = (value or "").strip().lower()
    if raw in {"male", "m"}:
        return "male"
    if raw in {"female", "f"}:
        return "female"
    return ""


def opposite_gender(value: str) -> str:
    if value == "male":
        return "female"
    if value == "female":
        return "male"
    return ""


def opposite_gender_pool(selected: Profile, profiles: list[Profile]) -> list[Profile]:
    target_gender = opposite_gender(normalize_gender(getattr(selected, "gender", "")))
    if not target_gender:
        return []
    selected_name = selected.name.strip().lower()
    return [
        p
        for p in profiles
        if p.name.strip().lower() != selected_name and normalize_gender(getattr(p, "gender", "")) == target_gender
    ]


def profile_to_dict(profile: Profile) -> dict:
    return {
        "name": profile.name,
        "gender": normalize_gender(getattr(profile, "gender", "")) or "-",
        "age": profile.age,
        "city": profile.city,
        "likes": ", ".join(sorted(profile.likes)),
        "bio": profile.bio,
        "photo_url": getattr(profile, "photo_url", ""),
    }


def profile_table_row(profile: Profile) -> dict:
    return {
        "name": profile.name,
        "gender": normalize_gender(getattr(profile, "gender", "")) or "-",
        "age": profile.age,
        "city": profile.city,
        "likes": ", ".join(sorted(profile.likes)),
        "bio": profile.bio,
        "photo": "Yes" if str(getattr(profile, "photo_url", "")).strip() else "No",
    }


def fallback_face(name: str) -> str:
    seed = (name or "guest").strip().replace(" ", "")
    return f"https://i.pravatar.cc/420?u={seed}"


def profile_photo(profile: Profile) -> str:
    photo = str(getattr(profile, "photo_url", "")).strip()
    if photo:
        return photo
    return fallback_face(profile.name)


def render_public_match_card(profile: Profile, rank: int, score: float) -> None:
    photo = profile_photo(profile)
    st.markdown(
        f"""
        <div class="public-match-card">
            <img src="{photo}" alt="{profile.name}" style="width: 100%; border-radius: 12px; margin-bottom: 8px;" />
            <div style="display: flex; justify-content: space-between; align-items: center; gap: 8px;">
                <strong>#{rank} {profile.name}</strong>
                <span class="score-pill">{score * 100:.1f}% match</span>
            </div>
        </div>
        """,
        unsafe_allow_html=True,
    )


def render_profile_gallery(profiles: list[Profile], *, columns_count: int = 4) -> None:
    if not profiles:
        return
    cols = st.columns(columns_count)
    for idx, profile in enumerate(profiles):
        with cols[idx % columns_count]:
            st.markdown('<div class="gallery-card">', unsafe_allow_html=True)
            st.image(profile_photo(profile), use_container_width=True)
            st.markdown(f'<div class="gallery-title">{profile.name}</div>', unsafe_allow_html=True)
            age_label = str(profile.age) if profile.age is not None else "-"
            city_label = profile.city or "-"
            gender_label = normalize_gender(getattr(profile, "gender", "")) or "-"
            st.markdown(f'<div class="gallery-meta">{gender_label.title()} | Age {age_label} | {city_label}</div>', unsafe_allow_html=True)
            bio = (profile.bio or "No bio added.").strip()
            if len(bio) > 80:
                bio = f"{bio[:77]}..."
            st.markdown(f'<div class="gallery-bio">{bio}</div>', unsafe_allow_html=True)
            st.markdown('</div>', unsafe_allow_html=True)


def make_profile(
    name: str,
    age: int | None,
    gender: str,
    city: str,
    likes: set[str],
    bio: str = "",
    photo_url: str = "",
) -> Profile:
    try:
        return Profile(
            name=name,
            age=age,
            gender=normalize_gender(gender),
            city=city,
            likes=likes,
            bio=bio,
            photo_url=photo_url,
        )
    except TypeError:
        profile = Profile(
            name=name,
            age=age,
            gender=normalize_gender(gender),
            city=city,
            likes=likes,
            bio=bio,
        )
        setattr(profile, "photo_url", photo_url)
        return profile


def normalize_columns(df: pd.DataFrame) -> pd.DataFrame:
    df = df.copy()
    df.columns = [str(col).strip().lower() for col in df.columns]
    return df


def parse_profiles_from_dataframe(df: pd.DataFrame) -> list[Profile]:
    df = normalize_columns(df)

    missing = {"name", "likes"} - set(df.columns)
    if missing:
        raise ValueError(f"Missing required columns: {', '.join(sorted(missing))}")

    profiles: list[Profile] = []
    for _, row in df.iterrows():
        name = str(row.get("name", "")).strip()
        if not name:
            continue

        raw_age = row.get("age")
        age = None
        if pd.notna(raw_age):
            try:
                age = int(raw_age)
            except (TypeError, ValueError):
                age = None

        city = ""
        raw_city = row.get("city")
        if pd.notna(raw_city):
            city = str(raw_city).strip()

        gender = ""
        raw_gender = row.get("gender")
        if pd.notna(raw_gender):
            gender = str(raw_gender).strip()

        bio = ""
        raw_bio = row.get("bio")
        if pd.notna(raw_bio):
            bio = str(raw_bio).strip()

        photo_url = ""
        raw_photo = row.get("photo_url")
        if pd.notna(raw_photo):
            photo_url = str(raw_photo).strip()

        likes = parse_likes(str(row.get("likes", "")))

        profiles.append(
            make_profile(
                name=name,
                age=age,
                gender=gender,
                city=city,
                likes=likes,
                bio=bio,
                photo_url=photo_url,
            )
        )

    return profiles


def parse_profiles_from_docx(file_bytes: bytes) -> list[Profile]:
    doc = Document(io.BytesIO(file_bytes))

    # Prefer table format when available.
    if doc.tables:
        table = doc.tables[0]
        rows = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        if len(rows) >= 2:
            header = rows[0]
            body = rows[1:]
            df = pd.DataFrame(body, columns=header)
            return parse_profiles_from_dataframe(df)

    # Fallback: parse paragraph key-value pairs.
    profiles = []
    current: dict[str, str] = {}

    def flush_profile() -> None:
        nonlocal current
        if not current.get("name"):
            current = {}
            return

        raw_age = current.get("age")
        age = None
        if raw_age:
            try:
                age = int(raw_age)
            except ValueError:
                age = None

        profiles.append(
            make_profile(
                name=current.get("name", "").strip(),
                age=age,
                gender=current.get("gender", "").strip(),
                city=current.get("city", "").strip(),
                likes=parse_likes(current.get("likes", "")),
                bio=current.get("bio", "").strip(),
                photo_url=current.get("photo_url", "").strip(),
            )
        )
        current = {}

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            flush_profile()
            continue

        if ":" in text:
            key, value = text.split(":", 1)
            current[key.strip().lower()] = value.strip()

    flush_profile()
    return profiles


def load_profiles(uploaded_file) -> list[Profile]:
    if uploaded_file is None:
        df = pd.read_csv(DUMMY_DATA_PATH)
        return parse_profiles_from_dataframe(df)

    suffix = Path(uploaded_file.name).suffix.lower()
    if suffix == ".csv":
        df = pd.read_csv(uploaded_file)
        return parse_profiles_from_dataframe(df)
    if suffix in {".xlsx", ".xls"}:
        df = pd.read_excel(uploaded_file)
        return parse_profiles_from_dataframe(df)
    if suffix == ".docx":
        return parse_profiles_from_docx(uploaded_file.read())

    raise ValueError("Unsupported file type. Please upload CSV, Excel, or DOCX.")


def main() -> None:
    st.set_page_config(page_title="Dating Match MVP", page_icon="💘", layout="wide")
    inject_styles()

    st.markdown(
        """
        <div class="hero">
            <h1 style="margin: 0 0 6px 0;">Dating Match MVP</h1>
            <p style="margin: 0; font-size: 1rem;">
                Load profile data, select a person, and instantly see the best compatibility matches.
            </p>
        </div>
        """,
        unsafe_allow_html=True,
    )

    with st.sidebar:
        st.subheader("Admin Data Upload")
        uploaded_file = st.file_uploader(
            "Upload profiles file",
            type=["csv", "xlsx", "xls", "docx"],
            help="Admin upload. If empty, dummy data is used.",
        )
        top_k = st.slider("Number of matches", min_value=1, max_value=10, value=5)

    try:
        base_profiles = load_profiles(uploaded_file)
    except Exception as exc:  # noqa: BLE001
        st.error(f"Failed to load profiles: {exc}")
        st.stop()

    if not base_profiles:
        st.warning("No valid base profiles found. Please check your upload.")
        st.stop()

    source_key = "dummy-source"
    if uploaded_file is not None:
        source_key = f"{uploaded_file.name}:{uploaded_file.size}"

    reset_profiles_in_memory(base_profiles, source_key)
    all_profiles = combined_profiles()

    st.markdown("### MVP Role Switch")
    view_mode = st.radio(
        "Choose portal to demo",
        ["Admin View", "Public View"],
        horizontal=True,
    )

    if view_mode == "Admin View":
        st.subheader("Admin Portal")
        st.caption("Full access: profile details, uploaded data, public submissions, and full match insights.")

        snapshot_left, snapshot_right = st.columns(2)
        with snapshot_left:
            st.markdown(
                f"""
                <div class="profile-card">
                    <strong>Total profiles in system:</strong> {len(all_profiles)}<br/>
                    <strong>Admin base profiles:</strong> {len(st.session_state.get('base_profiles', []))}<br/>
                    <strong>Public submissions:</strong> {len(st.session_state.get('public_submissions', []))}
                </div>
                """,
                unsafe_allow_html=True,
            )

        with snapshot_right:
            st.markdown(
                f"""
                <div class="profile-card">
                    <strong>Source:</strong> {'Uploaded admin file' if uploaded_file is not None else 'Dummy dataset'}<br/>
                    <strong>Top matches displayed:</strong> {top_k}
                </div>
                """,
                unsafe_allow_html=True,
            )

        st.markdown("#### Public Submissions (Full Info)")
        public_submissions = st.session_state.get("public_submissions", [])
        if not public_submissions:
            st.info("No public submissions yet.")
        else:
            render_profile_gallery(public_submissions, columns_count=3)
            with st.expander("View submissions table"):
                st.dataframe(pd.DataFrame([profile_table_row(p) for p in public_submissions]), use_container_width=True)

        st.markdown("#### Client Demo: Browse Singles")
        st.caption("Photo-first profile cards for stakeholder demos.")
        render_profile_gallery(all_profiles, columns_count=4)
        with st.expander("View all profiles admin table (optional)"):
            st.dataframe(pd.DataFrame([profile_table_row(p) for p in all_profiles]), use_container_width=True)

        if len(all_profiles) < 2:
            st.warning("Need at least 2 profiles to compute matches.")
            st.stop()

        names = sorted([p.name for p in all_profiles])
        st.markdown("### Select Profile Name")

        if "admin_selected_name" not in st.session_state or st.session_state["admin_selected_name"] not in names:
            st.session_state["admin_selected_name"] = names[0]

        current_idx = names.index(st.session_state["admin_selected_name"])
        nav_left, nav_mid, nav_right = st.columns([1, 2, 1])
        with nav_left:
            if st.button("Previous Profile", use_container_width=True):
                st.session_state["admin_selected_name"] = names[(current_idx - 1) % len(names)]
                st.rerun()
        with nav_mid:
            st.caption(f"Profile {current_idx + 1} of {len(names)}")
        with nav_right:
            if st.button("Next Profile", use_container_width=True):
                st.session_state["admin_selected_name"] = names[(current_idx + 1) % len(names)]
                st.rerun()

        selected_name = st.selectbox("Choose a profile", names, key="admin_selected_name")
        selected = pick_profile_by_name(all_profiles, selected_name)

        if selected is None:
            st.warning("Selected profile not found.")
            st.stop()

        st.markdown("#### Full Profile")
        profile_col, photo_col = st.columns([2, 1])
        with profile_col:
            st.markdown(
                f"""
                <div class="profile-card">
                    <h3 style="margin-top: 0;">{selected.name}</h3>
                    <p style="margin: 0 0 8px 0;"><strong>Gender:</strong> {normalize_gender(getattr(selected, 'gender', '')) or '-'} | <strong>Age:</strong> {selected.age if selected.age is not None else '-'} | <strong>City:</strong> {selected.city or '-'}</p>
                    <p style="margin: 0 0 8px 0;"><strong>Bio:</strong> {selected.bio or 'No bio added.'}</p>
                    <div>{profile_likes_html(selected.likes)}</div>
                </div>
                """,
                unsafe_allow_html=True,
            )
        with photo_col:
            st.image(profile_photo(selected), use_container_width=True)

        admin_candidates = opposite_gender_pool(selected, all_profiles)
        admin_results = top_matches(selected, admin_candidates, limit=top_k)
        st.markdown(f"#### Top {top_k} Matches (Admin Full View)")
        selected_gender = normalize_gender(getattr(selected, "gender", ""))
        if selected_gender not in {"male", "female"}:
            st.info("Set selected profile gender to Male or Female to compute matches.")
            st.stop()
        if not admin_candidates:
            st.info(f"No {opposite_gender(selected_gender).title()} profiles available for matching yet.")
            st.stop()
        for rank, (profile, score, common_likes) in enumerate(admin_results, start=1):
            st.markdown(
                f"""
                <div class="match-card">
                    <div style="display: flex; justify-content: space-between; align-items: center; gap: 16px;">
                        <div style="display: flex; align-items: center; gap: 12px; flex: 1;">
                            <img src="{profile_photo(profile)}" alt="{profile.name}" style="width: 78px; height: 78px; object-fit: cover; border-radius: 10px; border: 1px solid #f2d4c6;" />
                            <div>
                                <strong>#{rank} {profile.name}</strong><br/>
                                <span>{normalize_gender(getattr(profile, 'gender', '')).title()} | Age: {profile.age if profile.age is not None else '-'} | City: {profile.city or '-'}</span><br/>
                                <span>Bio: {profile.bio or '-'}</span>
                            </div>
                        </div>
                        <div class="match-score">{score * 100:.1f}%</div>
                    </div>
                    <div style="margin-top: 8px;"><strong>Common likes:</strong> {', '.join(sorted(common_likes)) if common_likes else 'None yet'}</div>
                </div>
                """,
                unsafe_allow_html=True,
            )

    else:
        st.subheader("Public Page")
        st.caption("Privacy mode: only match names, photos, and compatibility score are shown publicly.")

        with st.form("public_profile_form", clear_on_submit=True):
            public_left, public_right = st.columns(2)
            with public_left:
                public_name = st.text_input("Your name")
                public_gender = st.selectbox("Your gender", ["Female", "Male"])
                public_age = st.number_input("Your age", min_value=18, max_value=100, value=28, step=1)
                public_city = st.text_input("Your city")
            with public_right:
                public_likes = st.text_area("Your likes (comma-separated)", placeholder="coffee, travel, movies")
                public_bio = st.text_area("Short bio (optional)", placeholder="Optional")
                public_photo_url = st.text_input("Your photo URL (optional)", placeholder="https://...")

            public_submit = st.form_submit_button("Submit Profile")

        if public_submit:
            clean_name = public_name.strip()
            likes = parse_likes(public_likes)
            if not clean_name:
                st.error("Name is required.")
            elif not likes:
                st.error("Please provide at least one like.")
            elif profile_exists(combined_profiles(), clean_name):
                st.error("This name already exists. Please use another name.")
            else:
                new_public = make_profile(
                    name=clean_name,
                    age=int(public_age),
                    gender=public_gender,
                    city=public_city.strip(),
                    likes=likes,
                    bio=public_bio.strip(),
                    photo_url=public_photo_url.strip(),
                )
                st.session_state["public_submissions"] = list(st.session_state.get("public_submissions", [])) + [new_public]
                st.session_state["public_active_name"] = clean_name
                st.success("Profile submitted. Showing your top matches below.")
                st.rerun()

        all_profiles = combined_profiles()
        public_submissions = st.session_state.get("public_submissions", [])
        if not public_submissions:
            st.info("Submit a profile to see public match preview.")
            st.stop()

        public_names = sorted([p.name for p in public_submissions])
        default_name = st.session_state.get("public_active_name", public_names[0])
        if default_name not in public_names:
            default_name = public_names[0]
        default_idx = public_names.index(default_name)

        chosen_public_name = st.selectbox("Choose your submitted profile", public_names, index=default_idx)
        st.session_state["public_active_name"] = chosen_public_name
        selected_public = next((p for p in public_submissions if p.name == chosen_public_name), None)

        if selected_public is None:
            st.warning("Selected public profile not found.")
            st.stop()

        public_gender = normalize_gender(getattr(selected_public, "gender", ""))
        if public_gender not in {"male", "female"}:
            st.info("Set your gender to Male or Female to compute matches.")
            st.stop()

        public_candidates = opposite_gender_pool(selected_public, all_profiles)
        if not public_candidates:
            st.info(f"No {opposite_gender(public_gender).title()} profiles available for matching yet.")
            st.stop()

        public_results = top_matches(selected_public, public_candidates, limit=top_k)
        st.markdown(f"#### Top {top_k} Public Matches")
        st.caption("Only names, photos, and compatibility score are visible on this page.")

        if not public_results:
            st.info("No matches available yet.")
            st.stop()

        cols = st.columns(3)
        for idx, (profile, score, _) in enumerate(public_results, start=1):
            with cols[(idx - 1) % 3]:
                render_public_match_card(profile, idx, score)


if __name__ == "__main__":
    main()
